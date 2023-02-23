# Instrucciones de uso:
# Escribir en terminal 
# >python Copipaste.py {aqui va el nombre de la carpeta donde estan los archivos a modificar} 
# {nombre del archivo madre ya salvado (T)} {cantidad de archivos a modificar}

# ¿Qué hace?
# copia informacion del archivo madre a una cantidad de archivos que deseemos (sus nombres tiene una id
# de tal forma que los archivos modificados son subsiguientes al archivo madre 
# ej: 2031570.docx , 2031571.docx)

import docx
import docx2txt
import re
import numpy as np
from os import listdir, mkdir, remove
from os.path import isfile, join
import sys 

carpeta = sys.argv[1]
archivo_madre = sys.argv[2]
cantidad = sys.argv[3]

doc_madre = docx.Document(join(carpeta, archivo_madre+'.docx'))
table_madre = doc_madre.tables[0]

lista = [9,10,11,12,13,
         41,42,43,44,45,46,47,48,49,50,51,52,53,54,55,56,57,58,59,60,61]

print(int(cantidad))

cve_int = int(archivo_madre[:-1])

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def escribe_tabla(texto, x, tabla):
    tabla.rows[x].cells[2].add_paragraph(texto)
    delete_paragraph(tabla.rows[x].cells[2].paragraphs[0])

for i in range(int(cantidad)):
    archivo = str(cve_int+i+1)+'R'
    doc = docx.Document(join(carpeta, archivo+'.docx'))
    table = doc.tables[0]
    for j in lista:
        escribe_tabla(table_madre.rows[j].cells[2].text, j, table)
    doc.save(join(carpeta,archivo+'.docx'))
