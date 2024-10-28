# Instrucciones de uso:
# Escribir en terminal 
# >python Hace_pega.py {aqui va el nombre de la carpeta con los words que queremos convertir a txt} 
# {aqui nombre carpeta con archivos que deben presentar los mismos nombres de los archivos de la carpeta anterior}

# el resultado es una carpeta llamada "Words_a_txt"

import docx
import docx2txt
import re
import numpy as np
from os import listdir, mkdir, remove
from os.path import isfile, join
import sys 

dir_datos = sys.argv[1]
dir_form = sys.argv[2]

juzgados_lugar = np.load("Trabajo-Digitacion/Juzgados_lugar.npy")
juzgados_num = np.load("Trabajo-Digitacion/Juzgados_num.npy")
num_campo = {'TituloPublicacion': 5, 'TipoActuacion': 6, 'TipoConcesion': 7, 'Juzgado': 8, 'FechaPresenta': 9, 
                   'FechaSentencia': 12, 'Foja': 17, 
                   'VueltaInscripcion': 18, 'AnioInscripcion': 19, 'RegistroInscripcion': 20,
                   'NumeroInscripcion': 24, 'CausaRol': 25, 'RutRepresentante': 47, 'RutConcesionario': 57}

def get_tokens(text):
  tokens_espacio = text.split()    
  tokens = []


  for i in range(len(tokens_espacio)):
    if re.search('[\Wº]', tokens_espacio[i]) is None:
      tokens += [tokens_espacio[i]]
    else:
      tokens += re.split('([\Wº])', tokens_espacio[i])
  

  tokens = [i for i in tokens if i!='']

  return tokens

# Entrega una lista de los archivos en una carpeta
def filesInAFolder(folder):
    return [f for f in listdir(folder) if isfile(join(folder, f))]

# Convierte un archivo .docx a uno .txt 
def word_a_txt(archivo, dir):
    # Passing docx file to process function
    text = docx2txt.process(archivo)

    path = archivo.split("\\")
    archivo = path[-1]

    nom_completo = join(dir, archivo.replace(".docx",".txt"))

    # Saving content inside docx file into output.txt file
    with open(nom_completo, "w") as text_file:
        print(text, file=text_file)

nombre_archivos = filesInAFolder(dir_form)

nombre_carpeta_txt = "Words_a_txt"
mkdir(nombre_carpeta_txt)

for i in nombre_archivos:
    name_docxfile = join(dir_datos, i)
    word_a_txt(name_docxfile, nombre_carpeta_txt)



# entrega la posicion de un string en el texto, si no existe entrega False
def encuentra_string(string, tokens):
    for i in range(len(tokens)):
        if tokens[i] == string:
            return i
    return False

# dado un pattern entrega la posicion del 1er match en el texto, si no existe entrega False
# se le puede agregar la opcion ignora CASE
def encuentra_match(pattern, tokens, ignoraCase = False, desde=0):
    for i in range(desde,len(tokens)):
        if ignoraCase:
            if re.match(pattern, tokens[i], re.IGNORECASE):
                return i
        else:
            if re.match(pattern, tokens[i]):
                return i
    return False
              

# desde una posicion start (sin incluirla) junta los tokens hasta encontrar un string, 
# se puede colocar un stop a mano 
# separa las palabras con un espacio
def escribe_hasta(start, string, tokens, separador="", stop=10):
    j = start
    texto = ""
    while tokens[j] != string and j<start+stop:
        texto += (tokens[j]+separador)
        j+=1
    return texto

def escribe_mientras(start, pattern, tokens, separador="", stop=10):
    j = start
    texto = ""
    while re.match(pattern, tokens[j]) and j<start+stop:
        texto += (tokens[j]+separador)
        j+=1
    return texto

# Pasa un string que diga un numero ordinal a un int
def ordinal_to_int(string):
    ord = np.array(['Primer', 'Segundo', 'Tercer', 'Cuarto', 'Quinto', 
                    '1°', '2°', '3°','4°', '5°', '1er', '2do', '3er', '4to', '5to'])
    num = np.array([1, 2, 3, 4, 5, 1, 2, 3, 4, 5, 1, 2, 3, 4, 5])
    for i in range(len(ord)):
        if re.search(ord[i],string, re.IGNORECASE):
            return num[i]
    




def titulo_publicacion(tokens, output):
    hay_comillas = encuentra_string("“", tokens)
     
    if hay_comillas:
        TituloPublicacion = "SENTENCIA "+ escribe_hasta(hay_comillas+1,"”",tokens, " ")
        TituloPublicacion = TituloPublicacion[:-1]
        output['TituloPublicacion'] += [TituloPublicacion]
        return output

    hay_exp = encuentra_string("exploración", tokens)

    if hay_exp:
        TituloPublicacion = "SENTENCIA "+ escribe_hasta(hay_exp+1,",",tokens, " ")
        TituloPublicacion = TituloPublicacion[:-1]
        output['TituloPublicacion'] += [TituloPublicacion]
        return output

    hay_den = encuentra_string("denominada", tokens)

    if hay_den:
        TituloPublicacion = "SENTENCIA "+ escribe_hasta(hay_den+1,",",tokens, " ")
        TituloPublicacion = TituloPublicacion[:-1]
        output['TituloPublicacion'] += [TituloPublicacion]
        return output

    else:
        output['TituloPublicacion'] += [None]
        return output

def causa_rol(tokens, output):
    letra_V = encuentra_string("V", tokens)
    p = letra_V and tokens[letra_V+1]=="-" and re.match("^[0-9]{1,4}$",tokens[letra_V+2])
    q = tokens[letra_V+3]=="-" and re.match("^20[0-9]{2}$",tokens[letra_V+4])
    condicion_causa_rol = p and q
    if condicion_causa_rol:
        CausaRol = "V-"+tokens[letra_V+2]+"-"+tokens[letra_V+4]
        output['CausaRol'] += [CausaRol]
        return output
    else:
        output['CausaRol'] += [None]
        return output





def foja_num(tokens, output):
    hay_fojas = encuentra_match("fojas*", tokens, True)

    if hay_fojas:
        Foja = escribe_mientras(hay_fojas+1,"[.0-9]",tokens)
        output['Foja'] += [Foja]
        n = encuentra_match("[º°]", tokens, True, desde=hay_fojas)
        if n:
            output['NumeroInscripcion'] += [escribe_mientras(n+1,"[.0-9]",tokens)]
        else:
            output['NumeroInscripcion'] += [None]
        return output
       
    
    hay_fs = encuentra_match("fs", tokens, True)

    if hay_fs:
        Foja = escribe_mientras(hay_fs+2,"[.0-9]",tokens)
        output['Foja'] += [Foja]
        n = encuentra_match("[º°]", tokens, True, desde=hay_fs)
        if n:
            output['NumeroInscripcion'] += [escribe_mientras(n+1,"[.0-9]",tokens)]
        else:
            output['NumeroInscripcion'] += [None]
        return output
        

    else:
        output['Foja'] += [None]
        output['NumeroInscripcion'] += [None]
        return output

def num_inscripcion(tokens, output):
    n = encuentra_match("nº|número", tokens, True)
    if n:
        output['NumeroInscripcion'] += [escribe_mientras(n+1,"[.0-9]",tokens)]
        return output
    else:
        output['NumeroInscripcion'] += [None]
        return output



def juzgado(tokens, output):
    hay_juzgado = encuentra_match("juzgado", tokens, True)
    if hay_juzgado:
        Juzgado = "Juzgado de Letras de "
        x = tokens[hay_juzgado+4]+" "+tokens[hay_juzgado+5]+" "+tokens[hay_juzgado+6]+" "+tokens[hay_juzgado+7]
        for i in range(len(juzgados_lugar)):   
            if re.search(juzgados_lugar[i],x, re.IGNORECASE):
                if juzgados_num[i]:
                    Juzgado = Juzgado+juzgados_lugar[i]
                    ordinal = ordinal_to_int(tokens[hay_juzgado-1])
                    if ordinal:
                        if ordinal <= juzgados_num[i]:
                            Juzgado = str(ordinal)+" "+Juzgado
                else:
                    Juzgado = Juzgado+juzgados_lugar[i]
                break

        output['Juzgado'] += [Juzgado]
        return output
    else:
        output['Juzgado'] += [None]
        return output
    
def rut(texto, output):
    Ruts = re.findall("[0-9]{1,2}.[0-9]{3}.[0-9]{3}-[k0-9]", texto, re.I)
    if Ruts:
        dig = [int(r.split('.')[0]) for r in Ruts ]
        if len(dig) == 2:
            if dig[0]>25:
                output['RutConcesionario']+=[Ruts[0]]
                output['RutRepresentante']+=[Ruts[1]]
                return output
            if dig[1]>25:
                output['RutConcesionario']+=[Ruts[1]]
                output['RutRepresentante']+=[Ruts[0]]
                return output
            else:
                output['RutConcesionario']+=[None]
                output['RutRepresentante']+=[None]
                return output
        if len(dig) == 1:
            if dig[0]>25:
                output['RutConcesionario']+=[Ruts[0]]
                output['RutRepresentante']+=[None]
                return output
            else:
                output['RutConcesionario']+=[None]
                output['RutRepresentante']+=[None]
                return output
    output['RutConcesionario']+=[None]
    output['RutRepresentante']+=[None]
    return output

def vta(tokens, output):
    if encuentra_match("^vta$", tokens, True):
        output['VueltaInscripcion']+=['VTA']
        return output
    else:
        output['VueltaInscripcion']+=[None]
        return output

def analiza_archivo(file_name, output):
    f = open(file_name, "r")
    texto = f.read()
    tokens = get_tokens(texto)
    rut(texto, output)

    titulo_publicacion(tokens, output)
    juzgado(tokens, output)
    foja_num(tokens, output)
    vta(tokens,output)
    causa_rol(tokens, output)
    rut(texto,output)

    
    

    return output

def analiza_carpeta(dir_name, output):
    file_names = filesInAFolder(dir_name)
    for i in range(len(file_names)):
        analiza_archivo(join(dir_name,file_names[i]), output)
    return output


data_output = {'TituloPublicacion': [], 'Juzgado': [], 'FechaPresenta': [], 
                   'FechaSentencia': [], 'FechaInscripcion': [], 'Foja': [], 
                   'VueltaInscripcion': [], 'AnioInscripcion': [], 'RegistroInscripcion': [],
                   'NumeroInscripcion': [], 'CausaRol': [], 'RutRepresentante': [], 'RutConcesionario': []}


resultado = analiza_carpeta(nombre_carpeta_txt, data_output)

print(data_output)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


        

# escribe un texto en la fila x de una tabla
def escribe_tabla(texto, x, tabla):
    tabla.rows[x].cells[2].add_paragraph(texto)
    delete_paragraph(tabla.rows[x].cells[2].paragraphs[0])

def escribe_nuevocampo_tabla(x, tabla):
    tabla.rows[x].cells[1].add_paragraph('<FechaInscripcion>')
    delete_paragraph(tabla.rows[x].cells[1].paragraphs[0])

    tabla.rows[x].cells[3].add_paragraph('</FechaInscripcion>')
    delete_paragraph(tabla.rows[x].cells[3].paragraphs[0])

# agrega nueva fila despues de fila x de una tabla
def agrega_fila(x, tabla):
    tabla.add_row()
    insertion_row = tabla.rows[x]
    insertion_row._tr.addnext(tabla.rows[-1]._tr)


def punto_medio(tabla):
    if re.match('P',tabla.rows[84].cells[2].text):
        escribe_tabla(tabla.rows[85].cells[2].text, 116, tabla)
        escribe_tabla(tabla.rows[86].cells[2].text, 117, tabla)
        tabla.rows[87]._element.getparent().remove(tabla.rows[87]._element)
        tabla.rows[86]._element.getparent().remove(tabla.rows[86]._element)
        tabla.rows[85]._element.getparent().remove(tabla.rows[85]._element)
        tabla.rows[84]._element.getparent().remove(tabla.rows[84]._element)
        tabla.rows[83]._element.getparent().remove(tabla.rows[83]._element)
        tabla.rows[82]._element.getparent().remove(tabla.rows[82]._element)
        escribe_tabla('4', 81, tabla)
        escribe_tabla('1', 83, tabla)
        escribe_tabla('2', 89, tabla)
        escribe_tabla('3', 95, tabla)
        escribe_tabla('4', 101, tabla)

        


def escribe_carpeta(dir_name, data):
    file_names = filesInAFolder(dir_name)
    mkdir(dir_name+"PRE")
    for i in range(len(file_names)):
        locacion = join(dir_name,file_names[i])
        doc = docx.Document(join(dir_name,file_names[i]))
        table = doc.tables[0]
        escribe_tabla("SENTENCIA CONSTITUTIVA", num_campo['TipoActuacion'], table)
        escribe_tabla("EXPLORACION", num_campo['TipoConcesion'], table)
        escribe_tabla(data['TituloPublicacion'][i], num_campo['TituloPublicacion'], table)
        escribe_tabla(data['Juzgado'][i], num_campo['Juzgado'], table)
        escribe_tabla(data['Foja'][i], num_campo['Foja'], table)
        escribe_tabla(data['VueltaInscripcion'][i], num_campo['VueltaInscripcion'], table)
        escribe_tabla(data['NumeroInscripcion'][i], num_campo['NumeroInscripcion'], table)
        escribe_tabla(data['CausaRol'][i], num_campo['CausaRol'], table)
        #escribe_tabla(data['RutRepresentante'][i], num_campo['RutRepresentante'], table)
        #escribe_tabla(data['RutConcesionario'][i], num_campo['RutConcesionario'], table)
        punto_medio(table)
        agrega_fila(12, table)
        escribe_nuevocampo_tabla(13, table)
        nueva_locacion = join(dir_name+"PRE",file_names[i].replace(".docx", "R.docx"))
        doc.save(nueva_locacion)

escribe_carpeta(dir_form, data_output)